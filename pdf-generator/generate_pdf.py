#!/usr/bin/env python3
"""
PDF Generator - Mattias Thyr Standarddokument
Fyller i standardmall och genererar PDF
"""

import sys
import os
import subprocess
from datetime import datetime
from pathlib import Path
from docx import Document

class PDFGenerator:
    def __init__(self, template_path):
        self.template_path = template_path
        
    def create_document(self, title, content, output_name=None):
        """
        Skapar dokument från mall med given titel och innehåll
        
        Args:
            title: Dokumentets huvudtitel
            content: Textinnehåll (kan vara multiline string)
            output_name: Namn på utfil (utan extension)
        """
        if not output_name:
            output_name = f"dokument_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        # Läs mallen
        doc = Document(self.template_path)
        
        # Hitta och ersätt "Exempel på standarddokument" med användarens titel
        for paragraph in doc.paragraphs:
            if "Exempel på standarddokument" in paragraph.text:
                for run in paragraph.runs:
                    if "Exempel på standarddokument" in run.text:
                        run.text = run.text.replace("Exempel på standarddokument", title)
        
        # Ta bort allt innehåll efter titeln (från "Sammanfattning" och nedåt)
        # men behåll sidhuvud/sidfot
        start_removing = False
        paragraphs_to_remove = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            if "Sammanfattning" in paragraph.text or start_removing:
                start_removing = True
                paragraphs_to_remove.append(paragraph)
        
        # Ta bort paragrafer bakifrån för att inte påverka index
        for paragraph in paragraphs_to_remove:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
        
        # Lägg till nytt innehåll
        # Splitta innehåll på newlines och skapa stycken
        for line in content.split('\n'):
            if line.strip():  # Skippa tomma rader
                doc.add_paragraph(line.strip())
            else:
                doc.add_paragraph()  # Lägg till tom rad för spacing
        
        # Uppdatera datum i sidhuvud till dagens datum
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                for run in paragraph.runs:
                    if "2026-02-28" in run.text:
                        run.text = run.text.replace("2026-02-28", datetime.now().strftime('%Y-%m-%d'))
            
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if "2026-02-28" in paragraph.text:
                                for run in paragraph.runs:
                                    if "2026-02-28" in run.text:
                                        run.text = run.text.replace("2026-02-28", datetime.now().strftime('%Y-%m-%d'))
        
        # Spara DOCX
        output_docx = f"{output_name}.docx"
        doc.save(output_docx)
        print(f"✓ DOCX skapad: {output_docx}")
        
        # Konvertera till PDF med LibreOffice
        try:
            subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', '.', output_docx
            ], check=True, capture_output=True)
            
            output_pdf = f"{output_name}.pdf"
            print(f"✓ PDF skapad: {output_pdf}")
            print(f"\nFärdig! Filerna finns i: {os.getcwd()}")
            return output_pdf
            
        except subprocess.CalledProcessError as e:
            print(f"✗ Fel vid PDF-konvertering: {e}")
            print(f"  Men DOCX-filen är klar: {output_docx}")
            return None
        except FileNotFoundError:
            print("✗ LibreOffice saknas. Installera med:")
            print("  sudo apt install libreoffice-writer")
            print(f"  Men DOCX-filen är klar: {output_docx}")
            return None


def main():
    if len(sys.argv) < 2:
        print("PDF Generator - Mattias Thyr Standarddokument")
        print()
        print("Användning:")
        print("  python3 generate_pdf.py 'Titel' 'Innehåll' [filnamn]")
        print()
        print("Exempel:")
        print("  python3 generate_pdf.py 'Projektrapport' 'Detta är mitt innehåll...'")
        print("  python3 generate_pdf.py 'Rapport' 'Rad 1\nRad 2\nRad 3' min_rapport")
        print()
        print("Interaktivt läge:")
        print("  python3 generate_pdf.py -i")
        sys.exit(1)
    
    # Interaktivt läge
    if sys.argv[1] == '-i':
        print("=== PDF Generator - Interaktivt läge ===\n")
        title = input("Dokumenttitel: ")
        print("\nInnehåll (avsluta med Ctrl+D på tom rad):")
        lines = []
        try:
            while True:
                line = input()
                lines.append(line)
        except EOFError:
            pass
        content = '\n'.join(lines)
        output_name = input("\nFilnamn (utan .pdf): ").strip()
        if not output_name:
            output_name = None
    else:
        title = sys.argv[1]
        content = sys.argv[2] if len(sys.argv) > 2 else ""
        output_name = sys.argv[3] if len(sys.argv) > 3 else None
    
    # Hitta mallen
    script_dir = Path(__file__).parent
    template_path = script_dir.parent / "documents" / "Standarddokument.docx"
    
    if not template_path.exists():
        print(f"✗ Mall saknas: {template_path}")
        sys.exit(1)
    
    generator = PDFGenerator(str(template_path))
    generator.create_document(title, content, output_name)


if __name__ == "__main__":
    main()
