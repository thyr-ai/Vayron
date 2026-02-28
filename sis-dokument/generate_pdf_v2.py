#!/usr/bin/env python3
"""
SiS Dokument - Standarddokument Generator v2
Med valbara sidhuvud/sidfot-komponenter
"""

import sys
import os
import json
import subprocess
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class PDFGeneratorV2:
    def __init__(self):
        script_dir = Path(__file__).parent
        self.components_path = script_dir / "components.json"
        
        # Ladda komponenter
        with open(self.components_path, 'r', encoding='utf-8') as f:
            self.components = json.load(f)
    
    def show_menu(self, options, title):
        """Visa meny och returnera valt alternativ"""
        print(f"\n{title}")
        print("=" * len(title))
        for key, value in options.items():
            print(f"{key}. {value}")
        
        while True:
            choice = input("\nVälj nummer: ").strip()
            if choice in options:
                return choice, options[choice]
            print("Ogiltigt val, försök igen.")
    
    def get_custom_text(self, prompt):
        """Få anpassad text från användare"""
        return input(f"{prompt}: ").strip()
    
    def create_document(self, title, content, header_choice, footer_choice, 
                       custom_header=None, custom_footer=None, output_name=None):
        """Skapa dokument med valbara komponenter"""
        
        if not output_name:
            output_name = f"dokument_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        # Skapa nytt dokument
        doc = Document()
        
        # Sätt grundläggande stilar
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Georgia'
        font.size = Pt(11)
        font.color.rgb = RGBColor(0, 0, 0)  # Svart text
        
        # Lägg till sidhuvud
        section = doc.sections[0]
        header = section.header
        
        # Dagens datum
        today = datetime.now().strftime('%Y-%m-%d')
        
        # Skapa sidhuvud-tabell (2 kolumner)
        header_table = header.add_table(rows=1, cols=2, width=section.page_width)
        header_table.rows[0].cells[0].width = section.page_width // 2
        header_table.rows[0].cells[1].width = section.page_width // 2
        
        # Vänster kolumn (header-text)
        left_cell = header_table.rows[0].cells[0]
        left_para = left_cell.paragraphs[0]
        
        if header_choice == "8" and custom_header:
            header_text = custom_header
        elif header_choice == "1":
            header_text = ""  # Bara datum
        else:
            header_text = self.components['headers'][header_choice]
        
        if header_text:
            run = left_para.add_run(header_text)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(64, 64, 64)  # Mörkgrå
            run.font.name = 'Georgia'
        
        # Höger kolumn (datum)
        right_cell = header_table.rows[0].cells[1]
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = right_para.add_run(today)
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(64, 64, 64)  # Mörkgrå
        date_run.font.name = 'Georgia'
        
        # Lägg till horisontell linje under sidhuvud
        header.add_paragraph('_' * 80)
        
        # Lägg till titel
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.name = 'Georgia'
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc.add_paragraph()  # Tom rad
        
        # Lägg till innehåll
        for line in content.split('\n'):
            if line.strip():
                para = doc.add_paragraph(line.strip())
                para.style = 'Normal'
            else:
                doc.add_paragraph()
        
        # Lägg till sidfot
        footer = section.footer
        
        # Horisontell linje före sidfot
        footer.add_paragraph('_' * 80)
        
        footer_para = footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if footer_choice == "9" and custom_footer:
            footer_text = custom_footer
        elif footer_choice == "1":
            footer_text = ""  # Ingen sidfot
        elif footer_choice == "8":
            footer_text = "Sida "  # Sidnummer läggs till av Word
        else:
            footer_text = self.components['footers'][footer_choice]
        
        if footer_text:
            run = footer_para.add_run(footer_text)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(64, 64, 64)  # Mörkgrå
            run.font.name = 'Georgia'
        
        # Spara DOCX
        output_docx = f"{output_name}.docx"
        doc.save(output_docx)
        print(f"✓ DOCX skapad: {output_docx}")
        
        # Konvertera till PDF
        try:
            subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', '.', output_docx
            ], check=True, capture_output=True)
            
            output_pdf = f"{output_name}.pdf"
            print(f"✓ PDF skapad: {output_pdf}")
            print(f"\nFärdig! Filerna finns i: {os.getcwd()}")
            return output_pdf
            
        except subprocess.CalledProcessError as e:
            print(f"✗ Fel vid PDF-konvertering: {e}")
            print(f"  Men DOCX-filen är klar: {output_docx}")
            return None
        except FileNotFoundError:
            print("✗ LibreOffice saknas.")
            print(f"  Men DOCX-filen är klar: {output_docx}")
            return None


def main():
    print("=" * 60)
    print("SiS Dokument - Standarddokument Generator v2")
    print("=" * 60)
    
    gen = PDFGeneratorV2()
    
    # Steg 1: Titel
    title = input("\nDokumenttitel: ").strip()
    
    # Steg 2: Innehåll
    print("\nInnehåll (avsluta med Ctrl+D på tom rad):")
    lines = []
    try:
        while True:
            line = input()
            lines.append(line)
    except EOFError:
        pass
    content = '\n'.join(lines)
    
    # Steg 3: Välj sidhuvud
    header_choice, header_desc = gen.show_menu(gen.components['headers'], "Välj sidhuvud")
    custom_header = None
    if header_choice == "8":
        custom_header = gen.get_custom_text("Skriv anpassad sidhuvud-text")
    
    # Steg 4: Välj sidfot
    footer_choice, footer_desc = gen.show_menu(gen.components['footers'], "Välj sidfot")
    custom_footer = None
    if footer_choice == "9":
        custom_footer = gen.get_custom_text("Skriv anpassad sidfot-text")
    
    # Steg 5: Filnamn
    output_name = input("\nFilnamn (utan .pdf): ").strip()
    if not output_name:
        output_name = None
    
    # Skapa dokument
    gen.create_document(title, content, header_choice, footer_choice,
                       custom_header, custom_footer, output_name)


if __name__ == "__main__":
    main()
